import os
import streamlit as st
import PyPDF2
from langchain_openai.chat_models import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser, JsonOutputParser
from langchain.prompts import ChatPromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import DocArrayInMemorySearch
from langchain_openai.embeddings import OpenAIEmbeddings
from langchain.schema import Document
import json
from datetime import datetime, timedelta
import pandas as pd
import sqlite3
import hashlib
import jwt
from docx import Document as DocxDocument
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io
import plotly.express as px
import plotly.graph_objects as go
import os
import streamlit as st
import PyPDF2
from pathlib import Path
from langchain_openai.chat_models import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser, JsonOutputParser
from langchain.prompts import ChatPromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import DocArrayInMemorySearch
from langchain_openai.embeddings import OpenAIEmbeddings
from langchain.schema import Document
from pydantic import BaseModel  # Add this line
import json
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch

def ensure_directories():
    """Create necessary directories if they don't exist"""
    os.makedirs("src/static", exist_ok=True)

class LegalDatabaseManager:
    def __init__(self):
        self.project_root = Path(__file__).parent.parent.parent
        self.static_dir = self.project_root / "src" / "static"
        self.static_dir.mkdir(parents=True, exist_ok=True)
        
        # Define file paths
        self.codigo_penal_pdf = self.static_dir / "BOE-A-1995-25444-consolidado.pdf"
        self.lecrim_pdf = self.static_dir / "BOE-A-2000-323-consolidado.pdf"
        self.codigo_penal_txt = self.static_dir / "codigo_penal.txt"
        self.lecrim_txt = self.static_dir / "ley_enjuiciamiento_criminal.txt"

    def convert_pdf_to_txt(self, pdf_path: Path, txt_path: Path) -> None:
        """Convert a PDF file to TXT format."""
        try:
            with open(pdf_path, "rb") as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                
            with open(txt_path, "w", encoding="utf-8") as txt_file:
                txt_file.write(text)
            
            st.success(f"Archivo '{txt_path.name}' creado exitosamente.")
        except Exception as e:
            st.error(f"Error al convertir {pdf_path.name}: {str(e)}")

    def save_uploaded_file(self, uploaded_file, target_path: Path) -> bool:
        """Save an uploaded file to the specified path."""
        try:
            target_path.parent.mkdir(parents=True, exist_ok=True)
            with open(target_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            st.success(f"Archivo {target_path.name} cargado correctamente.")
            return True
        except Exception as e:
            st.error(f"Error al guardar {target_path.name}: {str(e)}")
            return False

    def check_and_setup_database(self):
        """Check if legal database files exist and set them up if needed."""
        
        # Check if PDF files exist
        if not self.codigo_penal_pdf.exists() or not self.lecrim_pdf.exists():
            st.warning("Los archivos PDF no se encuentran en el directorio. Por favor, vuelve a subirlos.")
            
            codigo_penal_pdf = st.file_uploader(
                "Sube el archivo PDF del Código Penal:", 
                type="pdf", 
                key="codigo_penal_pdf"
            )
            
            lecrim_pdf = st.file_uploader(
                "Sube el archivo PDF de la Ley de Enjuiciamiento Criminal:", 
                type="pdf", 
                key="ley_enjuiciamiento_pdf"
            )
            
            if codigo_penal_pdf:
                self.save_uploaded_file(codigo_penal_pdf, self.codigo_penal_pdf)
            
            if lecrim_pdf:
                self.save_uploaded_file(lecrim_pdf, self.lecrim_pdf)

        # Convert PDFs to TXT if needed
        if self.codigo_penal_pdf.exists() and not self.codigo_penal_txt.exists():
            self.convert_pdf_to_txt(self.codigo_penal_pdf, self.codigo_penal_txt)

        if self.lecrim_pdf.exists() and not self.lecrim_txt.exists():
            self.convert_pdf_to_txt(self.lecrim_pdf, self.lecrim_txt)

        # Final verification
        if self.codigo_penal_txt.exists() and self.lecrim_txt.exists():
            return True
        else:
            st.error("Por favor, asegúrate de que los archivos PDF se han convertido correctamente a .txt.")
            return False

    def get_codigo_penal_text(self) -> str:
        """Get the content of codigo_penal.txt."""
        try:
            if self.codigo_penal_txt.exists():
                with open(self.codigo_penal_txt, 'r', encoding='utf-8') as f:
                    return f.read()
            return ""
        except Exception as e:
            st.error(f"Error al leer codigo_penal.txt: {str(e)}")
            return ""

    def get_lecrim_text(self) -> str:
        """Get the content of ley_enjuiciamiento_criminal.txt."""
        try:
            if self.lecrim_txt.exists():
                with open(self.lecrim_txt, 'r', encoding='utf-8') as f:
                    return f.read()
            return ""
        except Exception as e:
            st.error(f"Error al leer ley_enjuiciamiento_criminal.txt: {str(e)}")
            return ""

# Template Management System
class TemplateManager:
    def __init__(self):
        self.templates = {
            "legal_analysis": """
            Analiza los siguientes aspectos legales:
            1. Fundamentos jurídicos aplicados
            2. Jurisprudencia relacionada
            3. Doctrina relevante
            4. Conclusiones principales
            
            Texto: {text}
            """,
            "evidence_analysis": """
            Analiza las pruebas presentadas:
            1. Pruebas documentales
            2. Pruebas testimoniales
            3. Pruebas periciales
            4. Valoración judicial
            
            Texto: {text}
            """,
            "procedural_review": """
            Revisa los aspectos procesales:
            1. Cumplimiento de plazos
            2. Requisitos formales
            3. Trámites realizados
            4. Recursos disponibles
            
            Texto: {text}
            """
        }
    
    def get_template(self, template_name):
        return self.templates.get(template_name)
    
    def add_template(self, name, content):
        self.templates[name] = content
    
    def list_templates(self):
        return list(self.templates.keys())

# Task Management System
class TaskManager:
    def __init__(self):
        self.tasks = []
    
    def add_task(self, task):
        self.tasks.append({
            "id": len(self.tasks) + 1,
            "description": task["description"],
            "deadline": task["deadline"],
            "status": "pending",
            "created_at": datetime.now(),
            "assigned_to": task.get("assigned_to"),
            "priority": task.get("priority", "medium")
        })
    
    def update_task_status(self, task_id, status):
        for task in self.tasks:
            if task["id"] == task_id:
                task["status"] = status
                break
    
    def get_pending_tasks(self):
        return [task for task in self.tasks if task["status"] == "pending"]
    
    def get_all_tasks(self):
        return self.tasks

# Database Manager
class DatabaseManager:
    def __init__(self, db_path="legal_analysis.db"):
        self.db_path = db_path
        self._create_tables()
    
    def _get_connection(self):
        """Create a new connection for each thread"""
        return sqlite3.connect(self.db_path)
    
    def _create_tables(self):
        """Initialize database tables"""
        conn = self._get_connection()
        try:
            cursor = conn.cursor()
            
            # Cases table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS cases (
                    id INTEGER PRIMARY KEY,
                    filename TEXT,
                    timestamp TEXT,
                    category TEXT,
                    summary TEXT
                )
            """)
            
            # Annotations table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS annotations (
                    id INTEGER PRIMARY KEY,
                    case_id INTEGER,
                    text TEXT,
                    annotation TEXT,
                    type TEXT,
                    timestamp TEXT,
                    FOREIGN KEY (case_id) REFERENCES cases(id)
                )
            """)
            
            # Tasks table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS tasks (
                    id INTEGER PRIMARY KEY,
                    case_id INTEGER,
                    description TEXT,
                    deadline TEXT,
                    status TEXT,
                    assigned_to TEXT,
                    priority TEXT,
                    created_at TEXT,
                    FOREIGN KEY (case_id) REFERENCES cases(id)
                )
            """)
            
            conn.commit()
        finally:
            conn.close()
    
    def save_case(self, case_data):
        """Save case data to database"""
        conn = self._get_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO cases (filename, timestamp, category, summary)
                VALUES (?, ?, ?, ?)
            """, (
                case_data["filename"],
                case_data["timestamp"],
                case_data["category"],
                case_data["summary"]
            ))
            conn.commit()
            return cursor.lastrowid
        finally:
            conn.close()
    
    def save_annotation(self, annotation_data):
        """Save annotation to database"""
        conn = self._get_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO annotations (case_id, text, annotation, type, timestamp)
                VALUES (?, ?, ?, ?, ?)
            """, (
                annotation_data["case_id"],
                annotation_data["text"],
                annotation_data["annotation"],
                annotation_data["type"],
                annotation_data["timestamp"]
            ))
            conn.commit()
        finally:
            conn.close()
    
    def get_case_annotations(self, case_id):
        """Get annotations for a specific case"""
        conn = self._get_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM annotations WHERE case_id = ?", (case_id,))
            return cursor.fetchall()
        finally:
            conn.close()
    
    def get_case(self, case_id):
        """Get case by ID"""
        conn = self._get_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM cases WHERE id = ?", (case_id,))
            return cursor.fetchone()
        finally:
            conn.close()
    
    def get_all_cases(self):
        """Get all cases"""
        conn = self._get_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM cases ORDER BY timestamp DESC")
            return cursor.fetchall()
        finally:
            conn.close()

# Security Manager
class SecurityManager:
    def __init__(self, secret_key):
        self.secret_key = secret_key
    
    def hash_password(self, password):
        return hashlib.sha256(password.encode()).hexdigest()
    
    def generate_token(self, user_id):
        return jwt.encode(
            {"user_id": user_id, "exp": datetime.utcnow() + timedelta(days=1)},
            self.secret_key,
            algorithm="HS256"
        )
    
    def verify_token(self, token):
        try:
            return jwt.decode(token, self.secret_key, algorithms=["HS256"])
        except jwt.ExpiredSignatureError:
            return None


class DocumentExporter:
    @staticmethod
    def export_to_word(content, filename):
        doc = DocxDocument()
        doc.add_heading("Análisis Legal", 0)
        doc.add_paragraph(content)
        
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io
    
    @staticmethod
    def export_to_pdf(content, filename):
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Crear estilos
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Centro
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#000080')  # Azul oscuro
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceBefore=6,
            spaceAfter=6,
            leading=14
        )
        
        # Construir el documento
        story = []
        
        # Título
        story.append(Paragraph("ANÁLISIS DE SENTENCIA JUDICIAL", title_style))
        story.append(Spacer(1, 12))
        
        # Convertir el contenido en secciones
        sections = content.split('\n\n')
        for section in sections:
            if section.strip():
                # Si la sección es un encabezado (en mayúsculas)
                if section.strip().isupper():
                    story.append(Paragraph(section, heading_style))
                else:
                    # Dividir en líneas para procesar cada una
                    lines = section.split('\n')
                    for line in lines:
                        if line.strip():
                            # Si la línea comienza con un título conocido
                            if any(line.startswith(prefix) for prefix in ['Órgano Judicial:', 'Fecha:', 'Nº Sentencia:', 'Tipo de Procedimiento:', 'Delito(s):', 'Artículos Aplicados:', 'Fallo:', 'Pena/Medida:']):
                                parts = line.split(':', 1)
                                if len(parts) == 2:
                                    story.append(Paragraph(
                                        f"<b>{parts[0]}:</b> {parts[1]}",
                                        normal_style
                                    ))
                            else:
                                story.append(Paragraph(line, normal_style))
                story.append(Spacer(1, 6))
        
        # Construir el PDF
        try:
            doc.build(story)
            buffer.seek(0)
            return buffer
        except Exception as e:
            print(f"Error al generar PDF: {str(e)}")
            raise
def init_session_state():
    if "api_key" not in st.session_state:
        st.session_state.api_key = ""
    if "pdf_file" not in st.session_state:
        st.session_state.pdf_file = None
    if "pdf_name" not in st.session_state:  # New
        st.session_state.pdf_name = None
    if "extracted_text" not in st.session_state:
        st.session_state.extracted_text = ""
    if "vectorstore_sentencias" not in st.session_state:
        st.session_state.vectorstore_sentencias = None
    if "vectorstore_leyes" not in st.session_state:
        st.session_state.vectorstore_leyes = None
    if "chain" not in st.session_state:
        st.session_state.chain = None
    if "embeddings" not in st.session_state:
        st.session_state.embeddings = None
    if "codigo_penal_docs" not in st.session_state:
        st.session_state.codigo_penal_docs = []
    if "context" not in st.session_state:
        st.session_state.context = None
    if "case_history" not in st.session_state:
        st.session_state.case_history = []
    if "saved_annotations" not in st.session_state:
        st.session_state.saved_annotations = {}
    if "document_summary" not in st.session_state:
        st.session_state.document_summary = None
    if "current_case_category" not in st.session_state:
        st.session_state.current_case_category = None
    if "analysis_complete" not in st.session_state:  # New
        st.session_state.analysis_complete = False
    if "template_manager" not in st.session_state:
        st.session_state.template_manager = TemplateManager()
    if "task_manager" not in st.session_state:
        st.session_state.task_manager = TaskManager()
    if "db_manager" not in st.session_state:
        st.session_state.db_manager = DatabaseManager()
    if "security_manager" not in st.session_state:
        st.session_state.security_manager = SecurityManager(secret_key="your-secret-key")
    if "current_case_id" not in st.session_state:
        st.session_state.current_case_id = None
    if "legal_db_manager" not in st.session_state:  # Add this
        st.session_state.legal_db_manager = LegalDatabaseManager()

# Core Functions
def load_legal_database(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        legal_text = file.read()
    return legal_text

def create_documents_from_text(text, source_name):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=20)
    chunks = text_splitter.split_text(text)
    documents = [Document(page_content=chunk, metadata={"source": source_name}) for chunk in chunks]
    return documents

def categorize_case(model, text):
    """Analysis for a single sentencing hearing"""
    categorization_template = """
    Analiza la siguiente sentencia judicial y proporciona un análisis estructurado en formato JSON con las siguientes claves exactas:
    {{
        "organo_judicial": "nombre del órgano judicial",
        "fecha": "fecha de la sentencia",
        "numero_sentencia": "número de la sentencia",
        "tipo_procedimiento": "tipo de procedimiento",
        "delitos": "lista de delitos",
        "articulos_aplicados": "artículos del Código Penal aplicados",
        "fundamentos_juridicos": "fundamentos jurídicos principales",
        "fallo": "decisión principal",
        "pena": "pena o medida impuesta",
        "resumen": "resumen del caso"
    }}

    Asegúrate de incluir todas las claves mencionadas en el resultado y que los valores sean strings.

    Sentencia: {text}
    """
    
    categorization_prompt = ChatPromptTemplate.from_template(categorization_template)
    
    # Define the expected JSON structure
    class CategorySchema(BaseModel):
        organo_judicial: str
        fecha: str
        numero_sentencia: str
        tipo_procedimiento: str
        delitos: str
        articulos_aplicados: str
        fundamentos_juridicos: str
        fallo: str
        pena: str
        resumen: str
    
    # Create a parser that will validate the structure
    json_parser = JsonOutputParser(pydantic_object=CategorySchema)
    
    categorization_chain = (
        categorization_prompt
        | model
        | json_parser
    )
    
    return categorization_chain.invoke({"text": text})
    # Define the expected JSON structure
    class CategorySchema(BaseModel):
        organo_judicial: str
        fecha: str
        numero_sentencia: str
        tipo_procedimiento: str
        delitos: str
        articulos_aplicados: str
        fundamentos_juridicos: str
        fallo: str
        pena: str
        resumen: str
    
    # Create a parser that will validate the structure
    json_parser = JsonOutputParser(pydantic_object=CategorySchema)
    
    categorization_chain = (
        categorization_prompt
        | model
        | json_parser
    )
    
    return categorization_chain.invoke({"text": text})

def process_document():
    if not st.session_state.pdf_file or not st.session_state.api_key:
        st.error("Por favor, proporciona un archivo PDF de sentencia válido y tu OpenAI API key.")
        return False
    
    try:
        with st.spinner("Analizando sentencia..."):
            # First ensure legal database is ready
            if not st.session_state.legal_db_manager.check_and_setup_database():
                st.error("Por favor, configura primero la base de datos legal.")
                return False

            # Configure OpenAI
            os.environ['OPENAI_API_KEY'] = st.session_state.api_key
            model = ChatOpenAI(openai_api_key=st.session_state.api_key, model="gpt-3.5-turbo")
            
            # Extract text from PDF
            pdf_reader = PyPDF2.PdfReader(st.session_state.pdf_file)
            extracted_text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if isinstance(page_text, str):
                    extracted_text += page_text
            
            st.session_state.extracted_text = extracted_text
            
            # Perform case categorization
            st.session_state.current_case_category = categorize_case(model, extracted_text)
            
            if not isinstance(st.session_state.case_history, list):
                st.session_state.case_history = []
            
            st.session_state.case_history.append({
                "filename": st.session_state.pdf_file.name,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "category": st.session_state.current_case_category,
                "tipo_procedimiento": st.session_state.current_case_category["tipo_procedimiento"],
                "organo_judicial": st.session_state.current_case_category["organo_judicial"],
                "delitos": st.session_state.current_case_category["delitos"]
            })
            
            # Create embeddings and vectorstore
            embeddings = OpenAIEmbeddings()
            sentencia_documents = create_documents_from_text(extracted_text, source_name="sentencia")
            st.session_state.vectorstore_sentencias = DocArrayInMemorySearch.from_documents(
                sentencia_documents, embeddings
            )
            
            # Load legal databases if not already loaded
            if not st.session_state.vectorstore_leyes:
                codigo_penal = st.session_state.legal_db_manager.get_codigo_penal_text()
                if codigo_penal:
                    st.session_state.codigo_penal_docs = create_documents_from_text(
                        codigo_penal, source_name="Código Penal"
                    )
                    st.session_state.vectorstore_leyes = DocArrayInMemorySearch.from_documents(
                        st.session_state.codigo_penal_docs, embeddings
                    )
                else:
                    st.warning("No se pudo cargar la base de datos legal. Algunas funciones pueden estar limitadas.")
            
            # Set up RAG context
            st.session_state.context = {
                "sentencias": st.session_state.vectorstore_sentencias.as_retriever(),
                "leyes": st.session_state.vectorstore_leyes.as_retriever() if st.session_state.vectorstore_leyes else None
            }
            
            # Save case to database
            case_data = {
                "filename": st.session_state.pdf_file.name,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "category": json.dumps(st.session_state.current_case_category),
                "summary": st.session_state.extracted_text[:500]
            }
            st.session_state.current_case_id = st.session_state.db_manager.save_case(case_data)
            
            # Create QA chain
            qa_template = """
            Basándote en la sentencia y el Código Penal, responde la siguiente pregunta 
            considerando todos los aspectos relevantes del caso.

            Sentencia: {sentencias}
            Código Penal: {leyes}

            Pregunta: {question}
            """
            qa_prompt = ChatPromptTemplate.from_template(qa_template)
            st.session_state.chain = (qa_prompt | model | StrOutputParser())
            
            return True
            
    except Exception as e:
        st.error(f"Error al procesar la sentencia: {str(e)}")
        return False

def save_annotation(text_selection, annotation_text, annotation_type):
    if st.session_state.current_case_id:
        annotation_data = {
            "case_id": st.session_state.current_case_id,
            "text": text_selection,
            "annotation": annotation_text,
            "type": annotation_type,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        st.session_state.db_manager.save_annotation(annotation_data)
        
        if text_selection not in st.session_state.saved_annotations:
            st.session_state.saved_annotations[text_selection] = []
        st.session_state.saved_annotations[text_selection].append({
            "annotation": annotation_text,
            "type": annotation_type,
            "timestamp": annotation_data["timestamp"]
        })

def export_document(content, format_type):
    if format_type == "docx":
        doc_io = DocumentExporter.export_to_word(content, "analisis_legal.docx")
        st.download_button(
            label="Descargar DOCX",
            data=doc_io,
            file_name="analisis_legal.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    elif format_type == "pdf":
        pdf_io = DocumentExporter.export_to_pdf(content, "analisis_legal.pdf")
        st.download_button(
            label="Descargar PDF",
            data=pdf_io,
            file_name="analisis_legal.pdf",
            mime="application/pdf"
        )

def get_answer(question):
    try:
        sentencias_relevant = st.session_state.context["sentencias"].get_relevant_documents(question)
        leyes_relevant = st.session_state.context["leyes"].get_relevant_documents(question)
        
        combined_sentencias = "\n".join([doc.page_content for doc in sentencias_relevant])
        combined_leyes = "\n".join([doc.page_content for doc in leyes_relevant])
        
        answer = st.session_state.chain.invoke({
            "sentencias": combined_sentencias,
            "leyes": combined_leyes,
            "question": question
        })
        
        return answer, sentencias_relevant, leyes_relevant
    
    except Exception as e:
        st.error(f"Error al procesar la pregunta: {str(e)}")
        return None, None, None

def show_home_page():
    st.header("Bienvenido al Sistema de Análisis Legal")
    
    # Dashboard summary
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Sentencias Analizadas", len(st.session_state.case_history))
    with col2:
        pending_tasks = len(st.session_state.task_manager.get_pending_tasks())
        st.metric("Tareas Pendientes", pending_tasks)
    with col3:
        annotation_count = sum(len(anns) for anns in st.session_state.saved_annotations.values())
        st.metric("Anotaciones Totales", annotation_count)
    
    # Recent activity
    st.subheader("Actividad Reciente")
    if st.session_state.case_history:
        recent_cases = pd.DataFrame(st.session_state.case_history[-5:])
        st.dataframe(recent_cases)
    else:
        st.info("No hay actividad reciente")

def show_document_analysis():
    st.header("Análisis de Sentencia")
    
    # File upload
    uploaded_file = st.file_uploader("Cargar Sentencia (PDF)", type="pdf", key="pdf_uploader")
    
    # Check if a new file was uploaded
    if not st.session_state.legal_db_manager.check_and_setup_database():
        st.warning("Por favor, configura primero la base de datos legal antes de analizar sentencias.")
        return
    if uploaded_file is not None and (st.session_state.pdf_name != uploaded_file.name):
        st.session_state.pdf_file = uploaded_file
        st.session_state.pdf_name = uploaded_file.name
        st.session_state.analysis_complete = False
        st.session_state.current_case_category = None
        st.session_state.extracted_text = ""
    
    # Show analyze button only if file is uploaded and not analyzed
    if st.session_state.pdf_file and not st.session_state.analysis_complete:
        if st.button("Analizar Sentencia"):
            if process_document():
                st.session_state.analysis_complete = True
                st.rerun()
    
    # Show analysis results if available
    if st.session_state.analysis_complete and st.session_state.current_case_category:
        st.success("Sentencia analizada correctamente")
        
        categorization = st.session_state.current_case_category
        
        # Display categorization in an organized format
        st.subheader("Análisis de la Sentencia")
        
        # Datos básicos
        st.write("**📋 Datos Básicos**")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.write("Órgano Judicial:")
            st.write(categorization["organo_judicial"])
        with col2:
            st.write("Fecha:")
            st.write(categorization["fecha"])
        with col3:
            st.write("Nº Sentencia:")
            st.write(categorization["numero_sentencia"])
        
        # Información procesal
        st.write("**⚖️ Información Procesal**")
        col1, col2 = st.columns(2)
        with col1:
            st.write("Tipo de Procedimiento:")
            st.write(categorization["tipo_procedimiento"])
        with col2:
            st.write("Delito(s):")
            st.write(categorization["delitos"])
        
        # Análisis legal
        st.write("**📚 Análisis Legal**")
        st.write("Artículos Aplicados:")
        st.write(categorization["articulos_aplicados"])
        
        st.write("Fundamentos Jurídicos:")
        st.write(categorization["fundamentos_juridicos"])
        
        # Decisión
        st.write("**🔨 Decisión**")
        col1, col2 = st.columns(2)
        with col1:
            st.write("Fallo:")
            st.write(categorization["fallo"])
        with col2:
            st.write("Pena/Medida:")
            st.write(categorization["pena"])
        
        # Resumen
        st.write("**📝 Resumen**")
        st.write(categorization["resumen"])
        
        # Debug info - Uncomment to debug
        # st.write("Debug - Categorization data:", categorization)
        
        # Q&A Interface
        st.subheader("Consultas sobre la Sentencia")
        question = st.text_input("Escriba su pregunta sobre la sentencia:", key="question_input")
        if question:
            if st.button("Obtener Respuesta", key="get_answer"):
                answer, sentencias, leyes = get_answer(question)
                if answer:
                    st.write("**Respuesta:**")
                    st.write(answer)
                    
                    with st.expander("Ver Referencias"):
                        if leyes:
                            st.subheader("Artículos Relacionados")
                            for doc in leyes:
                                st.write(doc.page_content)
        
        # Export options
        st.subheader("Exportar Análisis")
        if st.button("Exportar a PDF", key="export_pdf"):
            content = f"""
            ANÁLISIS DE SENTENCIA JUDICIAL
            
            Fecha de análisis: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            
            DATOS BÁSICOS
            Órgano Judicial: {categorization["organo_judicial"]}
            Fecha: {categorization["fecha"]}
            Nº Sentencia: {categorization["numero_sentencia"]}
            
            INFORMACIÓN PROCESAL
            Tipo de Procedimiento: {categorization["tipo_procedimiento"]}
            Delito(s): {categorization["delitos"]}
            
            ANÁLISIS LEGAL
            Artículos Aplicados: {categorization["articulos_aplicados"]}
            
            Fundamentos Jurídicos:
            {categorization["fundamentos_juridicos"]}
            
            DECISIÓN
            Fallo: {categorization["fallo"]}
            Pena/Medida: {categorization["pena"]}
            
            RESUMEN
            {categorization["resumen"]}
            """
            export_document(content, "pdf")
    elif st.session_state.pdf_file:
        st.info("Haz clic en 'Analizar Sentencia' para procesar el documento")
    else:
        st.info("Por favor, carga un archivo PDF con la sentencia a analizar")

def show_task_management():
    st.header("Gestión de Tareas")
    
    # Add new task
    with st.expander("Agregar Nueva Tarea"):
        task_description = st.text_area("Descripción de la tarea")
        col1, col2 = st.columns(2)
        with col1:
            deadline = st.date_input("Fecha límite")
        with col2:
            priority = st.selectbox("Prioridad", ["Alta", "Media", "Baja"])
        
        if st.button("Agregar Tarea"):
            st.session_state.task_manager.add_task({
                "description": task_description,
                "deadline": deadline.strftime("%Y-%m-%d"),
                "priority": priority
            })
            st.success("Tarea agregada correctamente")
    
    # Display tasks
    st.subheader("Tareas Pendientes")
    tasks = st.session_state.task_manager.get_pending_tasks()
    if tasks:
        for task in tasks:
            with st.container():
                col1, col2, col3 = st.columns([3, 1, 1])
                with col1:
                    st.write(task["description"])
                with col2:
                    st.write(f"Vence: {task['deadline']}")
                with col3:
                    if st.button("Completar", key=f"task_{task['id']}"):
                        st.session_state.task_manager.update_task_status(task["id"], "completed")
                        st.rerun()
    else:
        st.info("No hay tareas pendientes")

def show_annotations():
    st.header("Sistema de Anotaciones")
    
    if st.session_state.extracted_text and st.session_state.analysis_complete:
        # Text selection area
        selected_text = st.text_area("Seleccionar texto para anotar:", height=150)
        annotation_text = st.text_area("Escribir anotación:")
        annotation_type = st.selectbox(
            "Tipo de anotación:",
            ["Comentario Legal", "Precedente", "Acción Requerida", "Observación General"]
        )
        
        if st.button("Guardar Anotación"):
            save_annotation(selected_text, annotation_text, annotation_type)
            st.success("Anotación guardada correctamente")
        
        # Display annotations
        st.subheader("Anotaciones Guardadas")
        for text, annotations in st.session_state.saved_annotations.items():
            with st.expander(f"Texto: {text[:100]}..."):
                for ann in annotations:
                    st.write(f"**{ann['type']}**: {ann['annotation']}")
                    st.caption(f"Creado: {ann['timestamp']}")
    else:
        st.warning("Primero debes cargar y analizar una sentencia en la sección 'Análisis de Sentencia'")

def show_statistics():
    st.header("Estadísticas de Análisis")
    
    # Verificar si hay datos de casos analizados
    if st.session_state.current_case_category is not None:
        # Convertir el caso actual a DataFrame
        current_case = {
            "filename": st.session_state.pdf_name,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "tipo_procedimiento": st.session_state.current_case_category["tipo_procedimiento"],
            "organo_judicial": st.session_state.current_case_category["organo_judicial"],
            "delitos": st.session_state.current_case_category["delitos"]
        }
        
        # Basic statistics
        st.subheader("Resumen General")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Sentencias Analizadas", "1")
        with col2:
            annotation_count = len(st.session_state.saved_annotations.keys())
            st.metric("Anotaciones Realizadas", annotation_count)
        with col3:
            st.metric("Tareas Pendientes", 
                     len(st.session_state.task_manager.get_pending_tasks()))
        
        # Mostrar detalles del caso actual
        st.subheader("Detalles del Caso Actual")
        st.write("**Tipo de Procedimiento:**", current_case["tipo_procedimiento"])
        st.write("**Órgano Judicial:**", current_case["organo_judicial"])
        st.write("**Delitos:**", current_case["delitos"])
        
        # Visualizaciones
        st.subheader("Análisis del Caso")
        
        # Crear gráfico de los artículos aplicados
        articulos = st.session_state.current_case_category["articulos_aplicados"]
        if isinstance(articulos, str):  # Si es string, convertirlo a lista
            articulos = [art.strip() for art in articulos.replace('[', '').replace(']', '').split(',')]
        
        # Crear DataFrame para el gráfico
        df_articulos = pd.DataFrame({'Artículos': articulos})
        fig_articulos = px.bar(df_articulos['Artículos'].value_counts(), 
                             title='Artículos del Código Penal Aplicados')
        st.plotly_chart(fig_articulos)
        
        # Mostrar estadísticas de anotaciones si existen
        if st.session_state.saved_annotations:
            st.subheader("Análisis de Anotaciones")
            annotation_types = []
            for annotations in st.session_state.saved_annotations.values():
                for ann in annotations:
                    annotation_types.append(ann['type'])
            
            df_annotations = pd.DataFrame({'Tipo': annotation_types})
            fig_annotations = px.pie(df_annotations, 
                                   names='Tipo',
                                   title='Distribución de Tipos de Anotaciones')
            st.plotly_chart(fig_annotations)
        
        # Agregar timeline del caso
        st.subheader("Línea de Tiempo")
        timeline_data = pd.DataFrame([{
            'fecha': datetime.now(),
            'evento': 'Análisis de Sentencia'
        }])
        fig_timeline = px.line(timeline_data, 
                             x='fecha',
                             y=[1],
                             title='Actividad del Caso')
        st.plotly_chart(fig_timeline)
        
    else:
        st.info("No hay sentencias analizadas todavía")
        
    # Añadir botón para exportar estadísticas
    if st.session_state.current_case_category is not None:
        if st.button("Exportar Estadísticas"):
            stats_data = {
                "caso_actual": current_case,
                "anotaciones_total": annotation_count,
                "tareas_pendientes": len(st.session_state.task_manager.get_pending_tasks())
            }
            
            # Convertir a JSON y ofrecer para descarga
            json_stats = json.dumps(stats_data, indent=2)
            st.download_button(
                label="Descargar Estadísticas (JSON)",
                data=json_stats,
                file_name="estadisticas_analisis.json",
                mime="application/json"
            )

def show_templates():
    st.header("Gestión de Plantillas")
    
    # Add new template
    with st.expander("Agregar Nueva Plantilla"):
        template_name = st.text_input("Nombre de la plantilla")
        template_content = st.text_area("Contenido de la plantilla")
        
        if st.button("Guardar Plantilla"):
            st.session_state.template_manager.add_template(template_name, template_content)
            st.success("Plantilla guardada correctamente")
    
    # List existing templates
    st.subheader("Plantillas Disponibles")
    for template_name in st.session_state.template_manager.list_templates():
        with st.expander(template_name):
            st.code(st.session_state.template_manager.get_template(template_name))

def show_settings():
    st.header("Configuración")
    
    # General settings
    st.subheader("Configuración General")
    st.text_input("Ruta de base de datos", value="legal_analysis.db", disabled=True)
    st.text_input("Modelo OpenAI", value="gpt-3.5-turbo", disabled=True)
    
    # Export/Import settings
    st.subheader("Exportar/Importar Datos")
    if st.button("Exportar Todos los Datos"):
        data = {
            "case_history": st.session_state.case_history,
            "annotations": st.session_state.saved_annotations,
            "tasks": st.session_state.task_manager.tasks
        }
        st.download_button(
            "Descargar Datos",
            data=json.dumps(data, indent=2),
            file_name="legal_analysis_backup.json",
            mime="application/json"
        )

def main():
    # Initialize session state
    init_session_state()
    
    # Sidebar navigation
    st.sidebar.title("Sistema Legal RAG")
    navigation = st.sidebar.radio(
        "Navegación",
        ["Inicio", "Análisis de Sentencia", "Gestión de Tareas", "Anotaciones",
         "Estadísticas", "Plantillas", "Configuración"]
    )
    
    # Header with auth status
    with st.container():
        st.title("Sistema de Análisis Legal")
        
        # API Key input in sidebar
        st.sidebar.text_input(
            "OpenAI API Key",
            type="password",
            key="api_key",
            value=st.session_state.api_key
        )
    
    # Route to appropriate page
    if navigation == "Inicio":
        show_home_page()
    elif navigation == "Análisis de Sentencia":
        show_document_analysis()
    elif navigation == "Gestión de Tareas":
        show_task_management()
    elif navigation == "Anotaciones":
        show_annotations()
    elif navigation == "Estadísticas":
        show_statistics()
    elif navigation == "Plantillas":
        show_templates()
    elif navigation == "Configuración":
        show_settings()

if __name__ == "__main__":
    main()